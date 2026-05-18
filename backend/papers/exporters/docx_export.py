"""DOCX exporter — uses python-docx. Renders headings, paragraphs, tables,
inline markdown formatting (bold/italic/code/links), and bullet / numbered
lists.

This is a pragmatic markdown->docx converter that handles the constructs
the LangGraph writer actually emits. Anything fancier (nested lists,
images, footnotes) is rendered as plain text.
"""
from __future__ import annotations

import io
import re
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from db import PaperProject, PaperSection

_TABLE_LINE = re.compile(r"^\s*\|.+\|\s*$")
_TABLE_SEP = re.compile(r"^\s*\|?(\s*:?-{3,}:?\s*\|)+(\s*:?-{3,}:?\s*)\|?\s*$")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+\.\s+(.*)$")
_BLOCKQUOTE = re.compile(r"^\s*>\s?(.*)$")

# Inline tokens — order matters: bold (**) before italic (*).
_INLINE_PATTERN = re.compile(
    r"(\*\*([^*]+)\*\*"          # **bold**
    r"|__([^_]+)__"              # __bold__
    r"|\*([^*\n]+)\*"            # *italic*
    r"|_([^_\n]+)_"              # _italic_
    r"|`([^`\n]+)`"              # `code`
    r"|\[([^\]]+)\]\(([^)\s]+)\))"  # [text](url)
)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """python-docx has no high-level hyperlink API; build one via low-level
    OOXML. The text gets blue + underline so it visually reads as a link."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_run(paragraph, text: str, *, bold=False, italic=False, code=False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def _add_inline(paragraph, text: str) -> None:
    """Walk `text` and emit runs for each inline markdown token. Plain
    segments become regular runs; **bold** / *italic* / `code` / [link](url)
    get the matching formatting."""
    if not text:
        return
    cursor = 0
    for match in _INLINE_PATTERN.finditer(text):
        start, end = match.span()
        if start > cursor:
            _add_run(paragraph, text[cursor:start])
        bold_d, bold_u, ital_s, ital_u, code, link_text, link_url = match.groups()[1:]
        if bold_d is not None:
            _add_run(paragraph, bold_d, bold=True)
        elif bold_u is not None:
            _add_run(paragraph, bold_u, bold=True)
        elif ital_s is not None:
            _add_run(paragraph, ital_s, italic=True)
        elif ital_u is not None:
            _add_run(paragraph, ital_u, italic=True)
        elif code is not None:
            _add_run(paragraph, code, code=True)
        elif link_text is not None and link_url is not None:
            _add_hyperlink(paragraph, link_url, link_text)
        cursor = end
    if cursor < len(text):
        _add_run(paragraph, text[cursor:])


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            target = table.cell(r, c)
            target.text = ""  # clear default empty paragraph
            para = target.paragraphs[0]
            _add_inline(para, cell)
            # Header row → bold all runs
            if r == 0:
                for run in para.runs:
                    run.bold = True


def _consume_table(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows = []
    while i < len(lines) and _TABLE_LINE.match(lines[i]):
        if _TABLE_SEP.match(lines[i]):
            i += 1
            continue
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def _render_section_body(doc: Document, body: str) -> None:
    lines = (body or "").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Tables
        if _TABLE_LINE.match(line):
            rows, i = _consume_table(lines, i)
            _add_table(doc, rows)
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
            i += 1
            continue

        # Blank line
        if stripped == "":
            doc.add_paragraph("")
            i += 1
            continue

        # Bullet list
        m = _BULLET.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, m.group(1))
            i += 1
            continue

        # Numbered list
        m = _NUMBERED.match(line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, m.group(1))
            i += 1
            continue

        # Blockquote
        m = _BLOCKQUOTE.match(line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_inline(p, m.group(1))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Plain paragraph with inline formatting
        p = doc.add_paragraph()
        _add_inline(p, line)
        i += 1


def build_docx(project: PaperProject, sections: Iterable[PaperSection]) -> bytes:
    doc = Document()
    doc.add_heading(project.title or "Untitled Paper", level=0)

    meta_parts = []
    if project.domain:
        meta_parts.append(f"Domain: {project.domain}")
    if project.paper_type:
        meta_parts.append(f"Type: {project.paper_type}")
    if project.journal_type:
        meta_parts.append(f"Target journal: {project.journal_type}")
    if project.citation_style:
        meta_parts.append(f"Citation style: {project.citation_style.upper()}")
    if meta_parts:
        meta_para = doc.add_paragraph()
        run = meta_para.add_run(" · ".join(meta_parts))
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for section in sections:
        doc.add_heading(f"{section.order}. {section.title}", level=1)
        body = (section.body_md or "").strip() or "(section pending)"
        _render_section_body(doc, body)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
