"""Parse an uploaded research paper (PDF or DOCX) into a project + sections
shape compatible with the existing audit engine.

The audit engine reads simple attributes off each section
(``key``, ``title``, ``body_md``, optional ``guidance_json``) and off the
project (``citation_style``). We build lightweight ``SimpleNamespace``
stand-ins so we don't have to persist anything to the DB just to audit an
external paper.
"""
from __future__ import annotations

import io
import re
from types import SimpleNamespace
from typing import Any


# ---------------------------------------------------------------------------
# File text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(data: bytes) -> str:
    """Return the concatenated text of every page in the PDF."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "pypdf is not installed on the backend — add `pypdf` to requirements."
        ) from e
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            # Skip pages we can't decode — better partial text than failing the whole audit.
            continue
    return "\n\n".join(p.strip() for p in parts if p and p.strip())


def extract_text_from_docx(data: bytes) -> str:
    """Return paragraph-joined text from a DOCX file."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx is not installed on the backend.") from e
    doc = Document(io.BytesIO(data))
    paragraphs: list[str] = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            paragraphs.append(txt)
    # Pull table cell text too — Methodology / Results sections sometimes live in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    paragraphs.append(t)
    return "\n".join(paragraphs)


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch to the right extractor by file extension."""
    if not data:
        raise ValueError("Empty file")
    ext = ""
    if filename and "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(data)
    if ext in {"docx", "doc"}:
        return extract_text_from_docx(data)
    raise ValueError(f"Unsupported file type: .{ext or '?'} — please upload PDF or DOCX")


# ---------------------------------------------------------------------------
# Section splitting
# ---------------------------------------------------------------------------

# (canonical_key, display_title, alternation_group) — the alternation group is
# the regex source matched as the heading word itself. Order matters: more
# specific multi-word phrases come first so e.g. "Materials and Methods" wins
# over "Methods" when both could match.
_HEADING_DEFS: list[tuple[str, str, str]] = [
    ("title_authors", "Title, Authors & Affiliations",
        r"Title,?\s+Authors?(?:\s+(?:&|and)\s+Affiliations?)?"),
    ("abstract",      "Abstract",          r"Abstract"),
    ("keywords",      "Keywords",          r"Keywords?|Index\s+Terms"),
    ("introduction",  "Introduction",      r"Introduction|Background"),
    ("literature",    "Literature Review",
        r"Related\s+Work|Literature\s+Review|Prior\s+Work"),
    ("methodology",   "Methodology",
        r"Materials\s+and\s+Methods?|Proposed\s+(?:Method|Approach|System|Model|Framework)"
        r"|System\s+Design|Implementation|Methodology|Methods?|Approach"),
    ("experiments",   "Experiments",
        r"Experimental\s+Setup|Evaluation\s+Setup|Experiments?"),
    ("results",       "Results",
        r"Results\s+and\s+Discussion|Results|Findings|Evaluation"),
    ("discussion",    "Discussion",        r"Discussion"),
    ("conclusion",    "Conclusion",
        r"Conclusions?\s+and\s+Future\s+Work|Conclusions?|Summary"),
    ("future_work",   "Future Work",       r"Future\s+Work"),
    ("acknowledgments", "Acknowledgments", r"Acknowledg(?:e?ments?|ements?)"),
    ("references",    "References",        r"References|Bibliography|Works\s+Cited|Reference\s+List"),
]

# Single combined alternation, with a named group per definition so we can
# tell which heading matched. Anchored to paragraph start, allows the
# heading word to be followed by inline content on the same line — this is
# the bit that makes "2. Abstract The paradigm has…" detectable.
_HEADING_REGEX = re.compile(
    r"(?:^|\n)"                                           # paragraph break
    r"[ \t]*"                                             # leading whitespace
    r"(?:\#{1,6}[ \t]+)?"                                 # markdown ## prefix
    r"(?:\*\*[ \t]*)?"                                    # markdown bold open
    # Numbering prefix — allow up to 3 chained groups so headings like
    # "4. I. Introduction" (arabic + roman) and "10. VII. Conclusion" still
    # match. PDF extractors often emit this when the source had its own
    # roman numbering on top of section numbering.
    r"(?:(?:\d{1,2}(?:\.\d{1,2})*\.?|[IVX]{1,5}\.?)[ \t]+){0,3}"
    r"(?:" + "|".join(
        f"(?P<g{i}>{src})" for i, (_, _, src) in enumerate(_HEADING_DEFS)
    ) + r")"
    r"(?:[ \t]*\*\*)?"                                    # markdown bold close
    r"[ \t]*[:.\-]?",                                     # optional separator
    re.I,
)

def _guess_title(text: str) -> str:
    """Heuristic: paper title is usually the first non-empty, reasonably short
    line before the Abstract heading."""
    lines = text.split("\n")
    for ln in lines[:40]:
        # Strip leading markdown / numbering noise so titles like
        # "# Foo Bar" or "1. Foo Bar" come out clean.
        s = re.sub(r"^\s*(?:\#{1,6}\s+|\d+\.\s+|[IVX]+\.\s+)?", "", ln).strip(" *#")
        if not s:
            continue
        if re.search(r"@|\bUniversity\b|\bDept\b|\bDepartment\b", s):
            continue
        if re.fullmatch(r"[\d\W]+", s):
            continue
        if re.match(r"\s*Abstract\b", s, re.I):
            return "Uploaded Paper"
        words = s.split()
        if 3 <= len(words) <= 30:
            return s.rstrip(".")
    return "Uploaded Paper"


def _heading_for_match(match: re.Match) -> tuple[str, str] | None:
    """Return (canonical_key, display_title) for the first non-None named group.

    The display_title is the actual heading text the user wrote (e.g. "Related
    Work" or "Materials and Methods") — not the canonical hint — so audit
    cards reference what the user sees on the page rather than our internal
    bucket name. Falls back to the canonical disp if the match is somehow
    empty (shouldn't happen, but defensive).
    """
    for i, (key, disp, _) in enumerate(_HEADING_DEFS):
        actual = match.group(f"g{i}")
        if actual:
            # Title-case the matched text so "RELATED WORK" / "related work"
            # both normalise to "Related Work" without rewriting common-case
            # words mid-phrase ("and", "of", etc. stay lowercase for readers
            # who care, but our headings are short and title-case is fine).
            cleaned = " ".join(w for w in actual.split() if w).strip()
            return key, (cleaned or disp).title() if cleaned.isupper() or cleaned.islower() else cleaned or disp
    return None


def split_sections(full_text: str) -> tuple[str, list[dict[str, Any]]]:
    """Return (title, sections). Each section is a dict with key/title/body_md.

    Strategy: scan the whole text for every heading word that sits at a
    paragraph start (possibly prefixed with numbering or markdown markers).
    Each match becomes a section boundary; the body runs from the END of the
    heading match (so any inline content like "Abstract The paradigm has…"
    is captured) to the START of the next match.
    """
    text = full_text.replace("\r\n", "\n").replace("\r", "\n")
    title = _guess_title(text)

    hits: list[tuple[int, int, str, str]] = []  # (match_start, body_start, key, display)
    seen_keys: set[str] = set()
    for m in _HEADING_REGEX.finditer(text):
        info = _heading_for_match(m)
        if not info:
            continue
        key, disp = info
        if key in seen_keys:
            # Only the first occurrence of each canonical key — guards
            # against TOC + heading duplicates.
            continue
        seen_keys.add(key)
        hits.append((m.start(), m.end(), key, disp))

    sections: list[dict[str, Any]] = []
    for idx, (_match_start, body_start, key, disp) in enumerate(hits):
        body_end = hits[idx + 1][0] if idx + 1 < len(hits) else len(text)
        body = text[body_start:body_end].strip(" \t\n:.-")
        if not body:
            continue
        sections.append({"key": key, "title": disp, "body_md": body, "order": idx})

    # If we couldn't detect any sections, dump everything into a single "body"
    # section so the deterministic checks still have something to chew on.
    if not sections:
        sections = [{"key": "body", "title": "Body", "body_md": full_text.strip(), "order": 0}]

    return title, sections


# ---------------------------------------------------------------------------
# Project / section stand-ins for the audit engine
# ---------------------------------------------------------------------------

def build_project_and_sections(
    filename: str,
    data: bytes,
    paper_type: str,
    citation_style: str,
) -> tuple[Any, list[Any]]:
    """Top-level entry: file bytes → (project, sections) ready for run_audit."""
    full_text = extract_text(filename, data)
    if not full_text.strip():
        raise ValueError("Could not extract any text from the uploaded file.")

    title, section_dicts = split_sections(full_text)

    project = SimpleNamespace(
        id=0,
        title=title,
        domain="uploaded",
        paper_type=paper_type or "research_paper",
        paper_format=paper_type or "",
        citation_style=(citation_style or "ieee").lower(),
        journal_type="",
        num_sections=len(section_dicts),
    )

    sections: list[Any] = []
    for spec in section_dicts:
        sections.append(
            SimpleNamespace(
                id=0,
                key=spec["key"],
                title=spec["title"],
                order=spec.get("order", 0),
                body_md=spec["body_md"],
                guidance_json=None,
                word_count=len(spec["body_md"].split()),
                version=1,
            )
        )
    return project, sections
